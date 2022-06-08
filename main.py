import os
from docx import Document
from docx.shared import Inches
import pyttsx3

document = Document()

def speak(text):
    pyttsx3.speak(text)

def create_folder(directory):
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print("Creation of the directory %s failed" % directory)

# define the name of the directory to be created
path = create_folder("./images/")

print("Please insert your image on this current project in this folder: images")

is_profile_image_added = input("You have inserted the profile image? Yes or No: ")

if is_profile_image_added.lower() == 'yes':
    profile_image = input("Insert the image name of your profile image: ")
    document.add_picture("images/" + profile_image, width=Inches(2.0))
else:
    print("Ok..your cv will be created without profile image")

name = input("What is your name? ")
speak("Hello " + name + " How are you today?")
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

document.add_paragraph(
    name + ' | ' + phone_number + " | " + email
)

# About me
document.add_heading('About me')
about_me = input("Tell about yourself? ")
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From date: ")
to_date = input("To date: ")

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input("Describe your experience at " + company + ': ')
p.add_run(experience_details)

# More experience
while True:
    has_more_experiences = input('Do you have more experience? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From date: ")
        to_date = input("To date: ")

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input("Describe your experience at " + company + ': ')
        p.add_run(experience_details)
    else:
        break


# Skills
document.add_heading('Skills')
skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# More skills
while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "I consent to the processing of personal data in my curriculum vitae based on art. 13 of Legislative Decree 196/2003 and art. 13 GDPR 679/16"


document.save(name + "_cv.docx")